"""
DOCX Converter with Hyperlink Support

Converts markdown reports to DOCX format with proper hyperlinks.
"""

import os
import re
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.

    Args:
        paragraph: Docx paragraph object
        text: Display text for the link
        url: URL to link to
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a  w:r element)
    new_run = OxmlElement('w:r')

    # Set the run's style to the builtin hyperlink style
    rPr = OxmlElement('w:rPr')

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Add color (black for our B&W theme)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)

    new_run.append(rPr)

    # Join all the xml elements together
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)


def parse_markdown_links(text: str) -> List[Tuple[str, str, str]]:
    """
    Parse markdown links from text.

    Returns:
        List of tuples: (before_text, link_text, url, after_text)
    """
    # Pattern for markdown links: [text](url)
    pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    matches = []

    for match in re.finditer(pattern, text):
        matches.append({
            'start': match.start(),
            'end': match.end(),
            'text': match.group(1),
            'url': match.group(2)
        })

    return matches


def add_paragraph_with_links(doc: Document, text: str, style=None):
    """
    Add a paragraph that may contain markdown links, converting them to hyperlinks.

    Args:
        doc: Document object
        text: Text that may contain [text](url) links
        style: Paragraph style (optional)
    """
    paragraph = doc.add_paragraph(style=style)

    # Find all markdown links
    links = parse_markdown_links(text)

    if not links:
        # No links, just add plain text
        paragraph.add_run(text)
        return paragraph

    # Add text with links
    last_pos = 0
    for link in links:
        # Add text before the link
        if link['start'] > last_pos:
            paragraph.add_run(text[last_pos:link['start']])

        # Add the hyperlink
        add_hyperlink(paragraph, link['text'], link['url'])

        last_pos = link['end']

    # Add remaining text
    if last_pos < len(text):
        paragraph.add_run(text[last_pos:])

    return paragraph


def convert_markdown_to_docx(markdown_content: str, output_path: str) -> bool:
    """
    Convert markdown content to DOCX format with hyperlink support.

    Args:
        markdown_content: Markdown string to convert
        output_path: Path where DOCX should be saved

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Penetration Testing Report"
        doc.core_properties.subject = "Security Assessment"

        # Configure styles
        _configure_styles(doc)

        # Parse and convert markdown
        lines = markdown_content.split('\n')
        i = 0
        in_table = False
        table_rows = []
        in_code_block = False
        code_lines = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    in_code_block = False
                    _add_code_block(doc, '\n'.join(code_lines))
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in line and not line.strip().startswith('<!--'):
                if not in_table:
                    in_table = True
                    table_rows = []

                # Skip table separator line
                if re.match(r'^\s*\|[\s\-:]+\|\s*$', line):
                    i += 1
                    continue

                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(cells)

                i += 1
                continue
            else:
                if in_table:
                    # Finish table
                    _add_table(doc, table_rows)
                    in_table = False
                    table_rows = []

            # Handle headings
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                text = line.lstrip('#').strip()
                _add_heading(doc, text, level)
                i += 1
                continue

            # Handle horizontal rules
            if re.match(r'^\s*[-*_]{3,}\s*$', line):
                _add_horizontal_line(doc)
                i += 1
                continue

            # Handle bullet lists
            if line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:]
                # Check if this line has markdown links
                if '[' in text and '](' in text:
                    add_paragraph_with_links(doc, text, style='List Bullet')
                else:
                    _add_bullet_point(doc, text, 0)
                i += 1
                continue

            # Handle numbered lists
            if re.match(r'^\s*\d+\.\s+', line):
                text = re.sub(r'^\s*\d+\.\s+', '', line)
                _add_numbered_point(doc, text)
                i += 1
                continue

            # Handle paragraphs with potential links
            if line.strip() and not line.strip().startswith(('<', '>')):
                # Check if line contains markdown links
                if '[' in line and '](' in line:
                    add_paragraph_with_links(doc, line)
                else:
                    _add_paragraph_with_formatting(doc, line)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Default: regular paragraph
            if line.strip():
                doc.add_paragraph(line)

            i += 1

        # Save document
        doc.save(output_path)
        print(f"✓ DOCX conversion successful: {output_path}")
        return True

    except Exception as e:
        print(f"✗ DOCX conversion failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def _configure_styles(doc: Document):
    """Configure custom styles for the document."""
    styles = doc.styles

    # Configure Heading 1
    if 'Heading 1' in styles:
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(24)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 0)

    # Configure Heading 2
    if 'Heading 2' in styles:
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(18)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 0, 0)

    # Configure Heading 3
    if 'Heading 3' in styles:
        h3_style = styles['Heading 3']
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(0, 0, 0)


def _add_heading(doc: Document, text: str, level: int):
    """Add a heading to the document."""
    if level > 6:
        level = 6
    heading = doc.add_heading(text, level=level)


def _add_paragraph_with_formatting(doc: Document, text: str):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    paragraph = doc.add_paragraph()

    # Split by formatting markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue

        # Bold text
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic text
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Inline code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
        # Regular text
        else:
            paragraph.add_run(part)


def _add_bullet_point(doc: Document, text: str, level: int):
    """Add a bullet point to the document."""
    paragraph = doc.add_paragraph(text, style='List Bullet')


def _add_numbered_point(doc: Document, text: str):
    """Add a numbered point to the document."""
    paragraph = doc.add_paragraph(text, style='List Number')


def _add_table(doc: Document, rows: List[List[str]]):
    """Add a table to the document."""
    if not rows or len(rows) < 1:
        return

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Populate table
    for row_idx, row_data in enumerate(rows):
        table_row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = cell_data

            # Header row styling
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Set cell background color for header
                _set_cell_background(cell, "000000")

    # Add spacing after table
    doc.add_paragraph()


def _set_cell_background(cell, color: str):
    """Set background color for a table cell."""
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    except:
        pass  # If setting background fails, just skip it


def _add_code_block(doc: Document, code: str):
    """Add a code block to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Set background color (light gray)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.5)
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)


def _add_horizontal_line(doc: Document):
    """Add a horizontal line to the document."""
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(12)

    # Add a bottom border to simulate a horizontal line
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
